from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from typing import List, Optional
import io
import traceback
import logging

logger = logging.getLogger(__name__)

from ..database import get_db
from ..models import KnowledgeBase, BusinessLine
from ..schemas import KnowledgeBaseCreate, KnowledgeBaseUpdate, KnowledgeBaseOut

router = APIRouter(prefix="/api/knowledge", tags=["knowledge"])


def _to_out(item) -> KnowledgeBaseOut:
    return KnowledgeBaseOut(
        id=item.id,
        title=item.title,
        category=item.category,
        content=item.content,
        business_line_id=item.business_line_id,
        business_line_name=item.business_line.name if item.business_line else None,
        is_active=item.is_active,
        created_at=item.created_at,
        updated_at=item.updated_at,
    )


# --- File parsing helpers ---

def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"--- Página {i+1} ---\n{text.strip()}")
        return "\n\n".join(pages) if pages else "No se pudo extraer texto del PDF."
    except Exception as e:
        return f"Error leyendo PDF: {str(e)}"


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from Word .docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs) if paragraphs else "No se pudo extraer texto del documento."
    except Exception as e:
        return f"Error leyendo Word: {str(e)}"


def parse_excel(file_bytes: bytes, filename: str) -> str:
    """Extract text from Excel (.xlsx, .xls)."""
    try:
        import pandas as pd
        if filename.endswith('.csv'):
            df = pd.read_csv(io.BytesIO(file_bytes))
        else:
            xls = pd.ExcelFile(io.BytesIO(file_bytes))
            sheets = []
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                if not df.empty:
                    sheets.append(f"--- Hoja: {sheet_name} ---\n{df.to_string(index=False)}")
            return "\n\n".join(sheets) if sheets else "Archivo Excel vacío."
        return df.to_string(index=False) if not df.empty else "Archivo CSV vacío."
    except Exception as e:
        return f"Error leyendo Excel/CSV: {str(e)}"


def parse_text(file_bytes: bytes) -> str:
    """Read plain text file."""
    try:
        return file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        try:
            return file_bytes.decode('latin-1')
        except Exception:
            return "No se pudo leer el archivo de texto."


def parse_file(file_bytes: bytes, filename: str) -> str:
    """Auto-detect format and extract text."""
    fname = filename.lower()
    if fname.endswith('.pdf'):
        return parse_pdf(file_bytes)
    elif fname.endswith('.docx'):
        return parse_docx(file_bytes)
    elif fname.endswith(('.xlsx', '.xls')):
        return parse_excel(file_bytes, fname)
    elif fname.endswith('.csv'):
        return parse_excel(file_bytes, fname)
    elif fname.endswith(('.txt', '.md', '.json')):
        return parse_text(file_bytes)
    else:
        # Try as text
        return parse_text(file_bytes)


# --- Endpoints ---

@router.get("", response_model=List[KnowledgeBaseOut])
def get_all(category: Optional[str] = None, db: Session = Depends(get_db)):
    query = db.query(KnowledgeBase).order_by(KnowledgeBase.category, KnowledgeBase.title)
    if category:
        query = query.filter(KnowledgeBase.category == category)
    return [_to_out(item) for item in query.all()]


@router.get("/categories")
def get_categories():
    return [
        {"value": "productos", "label": "Productos y Medicamentos"},
        {"value": "protocolos", "label": "Protocolos de Visita"},
        {"value": "faq", "label": "Preguntas Frecuentes"},
        {"value": "general", "label": "Información General"},
        {"value": "archivo", "label": "Archivos Cargados"},
    ]


@router.post("", response_model=KnowledgeBaseOut)
def create(data: KnowledgeBaseCreate, db: Session = Depends(get_db)):
    item = KnowledgeBase(
        title=data.title,
        category=data.category,
        content=data.content,
        business_line_id=data.business_line_id,
        is_active=data.is_active if data.is_active is not None else True,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return _to_out(item)


@router.post("/upload")
def upload_file(
    file: UploadFile = File(...),
    category: str = Form("archivo"),
    business_line_id: Optional[str] = Form(None),
    db: Session = Depends(get_db),
):
    """Upload a file (PDF, Word, Excel, CSV, TXT) and extract content into knowledge base."""
    try:
        file_bytes = file.file.read()
        filename = file.filename or "archivo"

        # Parse the file
        content = parse_file(file_bytes, filename)

        if not content or content.startswith("Error") or content.startswith("No se pudo"):
            return {"success": False, "message": content, "entries_created": 0}

        # Split large content into chunks of ~3000 chars for better agent context
        chunks = []
        if len(content) > 4000:
            lines = content.split('\n')
            current_chunk = []
            current_len = 0
            chunk_num = 1
            for line in lines:
                if current_len + len(line) > 3000 and current_chunk:
                    chunks.append((chunk_num, '\n'.join(current_chunk)))
                    chunk_num += 1
                    current_chunk = [line]
                    current_len = len(line)
                else:
                    current_chunk.append(line)
                    current_len += len(line)
            if current_chunk:
                chunks.append((chunk_num, '\n'.join(current_chunk)))
        else:
            chunks = [(1, content)]

        bl_id = int(business_line_id) if business_line_id and business_line_id.strip() else None
        entries_created = 0

        for chunk_num, chunk_content in chunks:
            title = filename
            if len(chunks) > 1:
                title = f"{filename} (parte {chunk_num}/{len(chunks)})"

            item = KnowledgeBase(
                title=title,
                category=category,
                content=chunk_content,
                business_line_id=bl_id,
                is_active=True,
            )
            db.add(item)
            entries_created += 1

        db.commit()

        return {
            "success": True,
            "message": f"Archivo '{filename}' procesado exitosamente",
            "entries_created": entries_created,
            "total_characters": len(content),
            "filename": filename,
        }

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error procesando archivo: {str(e)}")


@router.post("/upload-multiple")
def upload_multiple_files(
    files: List[UploadFile] = File(...),
    category: str = Form("archivo"),
    business_line_id: Optional[str] = Form(None),
    db: Session = Depends(get_db),
):
    """Upload multiple files at once."""
    results = []
    total_entries = 0
    total_chars = 0
    bl_id = int(business_line_id) if business_line_id and business_line_id.strip() else None

    for file in files:
        try:
            file_bytes = file.file.read()
            filename = file.filename or "archivo"
            content = parse_file(file_bytes, filename)

            if not content or content.startswith("Error") or content.startswith("No se pudo"):
                results.append({"filename": filename, "success": False, "message": content})
                continue

            chunks = []
            if len(content) > 4000:
                lines = content.split('\n')
                current_chunk = []
                current_len = 0
                chunk_num = 1
                for line in lines:
                    if current_len + len(line) > 3000 and current_chunk:
                        chunks.append((chunk_num, '\n'.join(current_chunk)))
                        chunk_num += 1
                        current_chunk = [line]
                        current_len = len(line)
                    else:
                        current_chunk.append(line)
                        current_len += len(line)
                if current_chunk:
                    chunks.append((chunk_num, '\n'.join(current_chunk)))
            else:
                chunks = [(1, content)]

            entries = 0
            for chunk_num, chunk_content in chunks:
                title = filename
                if len(chunks) > 1:
                    title = f"{filename} (parte {chunk_num}/{len(chunks)})"
                item = KnowledgeBase(
                    title=title, category=category, content=chunk_content,
                    business_line_id=bl_id, is_active=True,
                )
                db.add(item)
                entries += 1

            total_entries += entries
            total_chars += len(content)
            results.append({"filename": filename, "success": True, "entries_created": entries, "characters": len(content)})
        except Exception as e:
            logger.error(f"Error processing {file.filename}: {e}")
            results.append({"filename": file.filename, "success": False, "message": str(e)})

    db.commit()
    success_count = sum(1 for r in results if r.get("success"))
    return {
        "success": success_count > 0,
        "message": f"{success_count} de {len(files)} archivos procesados exitosamente",
        "files": results,
        "total_entries_created": total_entries,
        "total_characters": total_chars,
    }


@router.put("/{item_id}", response_model=KnowledgeBaseOut)
def update(item_id: int, data: KnowledgeBaseUpdate, db: Session = Depends(get_db)):
    item = db.query(KnowledgeBase).filter(KnowledgeBase.id == item_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Entrada no encontrada")
    for key, value in data.model_dump(exclude_unset=True).items():
        setattr(item, key, value)
    db.commit()
    db.refresh(item)
    return _to_out(item)


@router.delete("/{item_id}")
def delete(item_id: int, db: Session = Depends(get_db)):
    item = db.query(KnowledgeBase).filter(KnowledgeBase.id == item_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Entrada no encontrada")
    db.delete(item)
    db.commit()
    return {"message": "Eliminado exitosamente"}
